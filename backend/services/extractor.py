import pdfplumber
import docx
import re
from pathlib import Path


def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text.strip()


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\@\/\+\#]', ' ', text)
    return text.strip()
